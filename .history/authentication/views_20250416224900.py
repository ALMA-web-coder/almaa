# import this to require login
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.contrib.auth.models import User
from django.contrib.auth import authenticate, login
# import this for sending email to user
from django.contrib.auth.tokens import default_token_generator
from django.contrib.sites.shortcuts import get_current_site
from django.core.mail import EmailMessage
from django.shortcuts import render, redirect
from django.template.loader import render_to_string
from django.utils.encoding import force_bytes
from django.utils.http import urlsafe_base64_encode, urlsafe_base64_decode
from datetime import datetime

from authentication.forms import UserRegistrationForm, UserLoginForm
from django.contrib.auth.models import Group
from .models import Application, Contact, Documents, WorkExperience, Other, Status
from django.shortcuts import render, get_object_or_404
from django import forms
from django.urls import reverse
from django.conf import settings
from .forms import (
    CertificateEducationalDocumentForm,
    DocumentUploadForm,
    EducationalBackgroundCertificateForm,
    EducationalBackgroundMastersForm,
    MastersEducationalDocumentForm,
    ProgramTypeForm,
    StatusForm,
    UserLoginForm,
    UserRegistrationForm,
    PersonalInfoForm,
    ContactAndAddressForm,
    WorkExperienceForm,
    OtherInfoForm,
    
)

from paynow import Paynow 
from django.utils.http import urlsafe_base64_encode
from django.utils.encoding import force_bytes
from django.template.loader import render_to_string
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
from openpyxl import Workbook
from io import BytesIO
from django.http import HttpResponse
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt

import base64
import json
# Create your views here.

@csrf_exempt  # Disable CSRF for verification from external payment service
def paynow_result(request):
    if request.method == 'POST':
        # Paynow will send payment result data as a POST request
        data = request.POST
        payment_id = data.get('payment_id')
        status = data.get('status')
        application_id = data.get('application_id')  # Assuming you send this

        # Find the application linked with the payment
        application = get_object_or_404(Application, id=application_id)
        
        if status == 'paid':
            application.is_paid = True
            application.save()
            messages.success(request, "Payment successful!")
        elif status in ['failed', 'pending']:
            application.is_paid = False
            application.save()
            messages.error(request, "Payment failed or is pending.")

        # Acknowledge receipt to Paynow
        return JsonResponse({'status': 'received'}, status=200)

    return JsonResponse({'error': 'Invalid request'}, status=400)

def application_list(request):
    applications = Application.objects.all()  # Fetch all applications
    return render(request, 'application/application_list.html', {'applications': applications})



def application_detail(request, application_id):
    application = get_object_or_404(Application, id=application_id)
    return render(request, 'application/application_detail.html', {'application': application})



@login_required(login_url='login')
def homepage(request):
    return render(request, 'homepage.html')


def register(request):
    form = UserRegistrationForm()

    if request.method == "POST":
        form = UserRegistrationForm(request.POST)
        if form.is_valid():
            user = form.save(commit=False)
            user.is_active = False
            user.save()

            # Assign user to the "Student" group
            try:
                student_group = Group.objects.get(name='Students')
                student_group.user_set.add(user)  # Add the user to the group
            except Group.DoesNotExist:
                # Handle the case if the group does not exist
                messages.error(request, 'User group not found. Please contact support.')

            # email user with activation link
            current_site = get_current_site(request)
            mail_subject = "Activate your account."

            # the message will render what is written in authentication/email_activation/activate_email_message.html
            message = render_to_string('authentication/email_activation/activate_email_message.html', {
                'user': form.cleaned_data['username'],
                'domain': current_site.domain,
                'uid': urlsafe_base64_encode(force_bytes(user.pk)),
                'token': default_token_generator.make_token(user),
            })
            to_email = form.cleaned_data['email']
            email = EmailMessage(
                mail_subject, message, to=[to_email]
            )
            email.send()
            messages.success(request, 'Account created successfully. Please check your email to activate your account.')
            return redirect('login')
        else:
            messages.error(request, 'Account creation failed. Please try again! Password should be at least 8 characters long and include a mix of letters, numbers, and special characters.')

    return render(request, 'authentication/register.html', {
        'form': form
    })
# to activate user from email
def activate(request, uidb64, token):
    try:
        uid = urlsafe_base64_decode(uidb64).decode()
        user = User.objects.get(pk=uid)
    except(TypeError, ValueError, OverflowError, User.DoesNotExist):
        user = None
    if user is not None and default_token_generator.check_token(user, token):
        user.is_active = True
        user.save()
        return render(request, 'authentication/email_activation/activation_successful.html')
    else:
        return render(request, 'authentication/email_activation/activation_unsuccessful.html')


def login_view(request):
    if request.method == 'POST':
        form = UserLoginForm(data=request.POST)
        if form.is_valid():
            username = form.cleaned_data.get('username')
            password = form.cleaned_data.get('password')
            user = authenticate(request, username=username, password=password)

            if user is not None:
                # Successful login
                login(request, user)  # Log the user in

                # Check user group and redirect accordingly
                if user.groups.filter(name='Admin').exists():
                    return redirect('/admin/')  # Redirect for Registrars
                elif user.groups.filter(name='Students').exists():
                    return redirect('homepage')  # Redirect for Students
                else:
                    # Redirect or message for users not in the specified groups
                    messages.warning(request, 'You do not have permission to access this application.')
                    return redirect('login')  # Redirect back to login or another page
            else:
                messages.error(request, 'Invalid username or password.')  # Invalid credentials
        else:
            # Handle form errors by sending feedback to the user
            for field, errors in form.errors.items():
                for error in errors:
                    messages.error(request, error)

    else:
        form = UserLoginForm()  # Initialize empty form for GET requests

    # Render the login template with the form
    return render(request, 'authentication/login.html', {'form': form})



def application_view(request):
    user = request.user
    step = request.GET.get('step', '1')
    
    # Handle the first step of the process: selecting application type
    if step == '1':
        return step_one_handler(request, user)

    # Retrieve the user's most recent application only for the next steps
    application = Application.objects.filter(user=user).last()

    # Ensure an application exists for further steps
    if application is None:
        return redirect(reverse('authentication:application_view') + '?step=1')

    # Map steps to their corresponding handlers
    step_handlers = {
        '2': step_two_handler,
        '3': step_three_handler,
        '4': step_four_handler,
        '5': step_five_handler,
        '6': step_six_handler,
        '7': step_seven_handler,
        '8': step_eight_handler,
        '9': payment_handler,
        '10': step_nine_handler,
    }

    # Call the appropriate handler based on the current step
    if step in step_handlers:
        return step_handlers[step](request, application)
    else:
        return redirect(reverse('application_view') + '?step=1')


def step_one_handler(request, user):
    form = ProgramTypeForm(request.POST or None)  # Use POST data if available
    if request.method == 'POST':
        if form.is_valid():
            application = Application(user=user)
            application.application_type = form.cleaned_data['application_type']
            application.save()
            return redirect(reverse('application_view') + '?step=2')
    return render(request, 'application/program_selection.html', {'form': form})

def step_two_handler(request, application):
    form = PersonalInfoForm(request.POST or None, instance=application)
    if request.method == 'POST' and form.is_valid():
        form.save()
        return redirect(reverse('application_view') + '?step=3')
    return render(request, 'application/personal.html', {'form': form})

def step_three_handler(request, application):
    documents, _ = Documents.objects.get_or_create(application=application)

    form = DocumentUploadForm(request.POST or None, request.FILES or None, instance=documents)
    if request.method == 'POST' and form.is_valid():
        documents = form.save()
        return redirect(reverse('application_view') + '?step=4')

    return render(request, 'application/personal_document.html', {'form': form})

def step_four_handler(request, application):
    contact, _ = Contact.objects.get_or_create(application=application)

    form = ContactAndAddressForm(request.POST or None, instance=contact)
    if request.method == 'POST' and form.is_valid():
        form.save()
        return redirect(reverse('application_view') + '?step=5')

    return render(request, 'application/contact_and_address.html', {'form': form})

def step_five_handler(request, application):
    form_class = EducationalBackgroundMastersForm if application.application_type == 'masters' else EducationalBackgroundCertificateForm
    
    # existing_instance = application.educational.first() if application.educational.exists() else None
    
    template_name = 'application/masters_educational_background.html' if application.application_type == 'masters' else 'application/certificate_educational_background.html'
    existing_instance = application.educational.first() if application.educational.exists() else None
        
    form = form_class(request.POST or None, instance=existing_instance)
    if request.method == 'POST' and form.is_valid():
        educational = form.save(commit=False)
        educational.application = application
        educational.save()
        return redirect(reverse('application_view') + '?step=6')

    print(form.errors)
    return render(request, template_name, {'form': form})

def step_six_handler(request, application):
    if application.application_type == 'masters':
        form_class = MastersEducationalDocumentForm 
        template_name = 'application/mastersdoc.html'
        existing_instance = application.masters.first() if application.masters.exists() else None
    else:  # Certificates
        form_class = CertificateEducationalDocumentForm  
        template_name = 'application/certificatedoc.html'
        existing_instance = application.certificate.first() if application.certificate.exists() else None

    form = form_class(request.POST or None, request.FILES or None, instance=existing_instance)
    if request.method == 'POST' and form.is_valid():
        educational = form.save(commit=False)
        educational.application = application
        educational.save()
        return redirect(reverse('application_view') + '?step=7')

    return render(request, template_name, {'form': form})

def step_seven_handler(request, application):
    WorkExperienceFormSet = forms.inlineformset_factory(Application, WorkExperience, form=WorkExperienceForm, extra=1, can_delete=False)
    formset = WorkExperienceFormSet(request.POST or None, queryset=application.work_experiences.all())
    if request.method == 'POST' and formset.is_valid():
        instances = formset.save(commit=False)
        for instance in instances:
            instance.application = application
            instance.save()
        return redirect(reverse('application_view') + '?step=8')

    return render(request, 'application/work_experience.html', {'formset': formset})

def step_eight_handler(request, application):
    other, _ = Other.objects.get_or_create(application=application)

    form = OtherInfoForm(request.POST or None, instance=other)
    if request.method == 'POST':
        if form.is_valid():
            other_info = form.save()  # No need to commit=False then save again.  form.save() handles everything

            return redirect(reverse('application_view') + '?step=9')
        else:
            print(form.errors)  # Print form errors to the console for debugging

    return render(request, 'application/other.html', {'form': form})
def payment_handler(request, application):
    if request.method == 'POST':
        # Here you should include your logic for forming the payment request
        amount = 0.00 # Fetch the amount depending on the application or fixed value
        # Create a payment link or payload here as required by PayNow
        paynow_url = "https://www.paynow.co.zw/alma"

        # You will usually need to send this data alongside the request, e.g., in JSON format.
        payment_data = {
            "amount": amount,  # Amount to be paid
            "currency": "USD",  # Currency type, adjust if necessary
            "application_id": application.id,  # The app ID for tracking purposes, if needed
            "email": request.user.email,  # User’s email typically passed along for payment notification
            # Add any additional data required by PayNow
        }

        # Optionally convert your data to JSON and encode if needed
        payload = base64.b64encode(json.dumps(payment_data).encode('utf-8')).decode('utf-8')

        # Redirect the user to PayNow with the payload data
        return redirect(f"{paynow_url}?data={payload}")

    return render(request, 'application/payment.html', {
        'application': application,
        'amount': 5000,  # Example amount, replace with the appropriate logic to retrieve it
    })
def step_nine_handler(request, application):
    return render(request, 'application/success.html', {
        'current_year': datetime.now().year,
        'email': request.user.email,
    })
    
    
def track_application(request):
    application = None
    latest_status = None
    error_message = None
    
    if request.method == 'POST':
        national_id = request.POST.get('national_id')
        first_name = request.POST.get('first_name')
        last_name = request.POST.get('last_name')
        
        try:
            
            application = Application.objects.get(
                national_id=national_id,
                first_name__iexact=first_name,
                last_name__iexact=last_name
            )
            
            latest_status = application.statuses.order_by('-id').first()  
        except Application.DoesNotExist:
            error_message = "No application found with the provided details."

    
    return render(request, 'application/track_application.html', {
        'application': application,
        'latest_status': latest_status,
        'error_message': error_message
    })
    
    
@login_required 
def update_status(request, application_id):
    application = get_object_or_404(Application, id=application_id)
    
    
    status, created = Status.objects.get_or_create(application=application)
    
    if request.method == 'POST':
        form = StatusForm(request.POST, instance=status)
        if form.is_valid():
            form.save()  
            return redirect('application_detail', application_id=application.id) 
    else:
        form = StatusForm(instance=status)

    return render(request, 'application/update_status.html', {'form': form, 'application': application})
   
def registrar_view(request, application_id):
    application = get_object_or_404(Application, id=application_id)
    
    status_instance = Status.objects.filter(application=application).order_by('-updated_at').first()

    if status_instance is None:
        
        status_instance = Status(application=application, status='Pending')
        status_instance.save()

    if request.method == "POST":
        form = StatusForm(request.POST, instance=status_instance)
        if form.is_valid():
            form.save() 
            return redirect('application_detail', application_id=application_id)  
    else:
        form = StatusForm(instance=status_instance)

    
    return render(request, 'application/application_detail.html', {
        'form': form,
        'application': application,
        'status': status_instance.status,  
        'last_update': status_instance.last_update  
    }) 
    
    
   
def download_pdf(request, application_id):
    application = get_object_or_404(Application, id=application_id)
    contact = application.contacts.first()  # Assuming there's at least one contact related
    masters_edu = application.educational.first()  # Assuming the first educational record for masters
    certificate_edu = application.certificate_edu.first()  # Assuming the first educational record for certificate
    work_experience = application.work_experiences.first()  # Fetch first work experience for simplicity
    documents = application.documents.first()  # Assuming one set of documents is related

    # Create a PDF document
    buffer = BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    
    p.drawString(100, height - 40, f"Application Detail for {application.first_name} {application.last_name}")
    p.drawString(100, height - 60, f"Application Type: {application.get_application_type_display()}")
    p.drawString(100, height - 80, f"First Name: {application.first_name}")
    p.drawString(100, height - 100, f"Last Name: {application.last_name}")
    p.drawString(100, height - 120, f"Date Of Birth: {application.date_of_birth}")
    p.drawString(100, height - 140, f"Gender: {application.gender}")
    p.drawString(100, height - 160, f"National ID: {application.national_id}")
    p.drawString(100, height - 180, f"Title: {application.get_title_display()}")
    p.drawString(100, height - 200, f"Marital Status: {application.get_marital_status_display()}")
    p.drawString(100, height - 220, f"Place of Birth: {application.place_of_birth}")
    p.drawString(100, height - 240, f"Citizenship: {application.citizenship}")
    p.drawString(100, height - 260, f"Country: {application.country}")
    p.drawString(100, height - 280, f"Disability: {application.disability}")
    p.drawString(100, height - 300, f"Disability Details: {application.disability_details}")
    
    if contact:
        p.drawString(100, height - 320, f"Email: {contact.email}")
        p.drawString(100, height - 340, f"Phone Number: {contact.phone_number}")
        p.drawString(100, height - 360, f"Permanent Address: {contact.permanent_address_no_street}")
        if contact.permanent_address_apt_unit:
            p.drawString(100, height - 380, f"Apt/Unit: {contact.permanent_address_apt_unit}")
        p.drawString(100, height - 400, f"State/Province: {contact.permanent_address_state_province}")

    if masters_edu:
        p.drawString(100, height - 420, f"Programme Applied For (Masters): {masters_edu.get_programme_applied_for_display()}")
        p.drawString(100, height - 440, f"Institution: {masters_edu.awarding_institution}")
        p.drawString(100, height - 460, f"Major Subjects: {masters_edu.major_subjects}")
    
    if certificate_edu:
        p.drawString(100, height - 480, f"Programme Applied For (Certificate): {certificate_edu.get_programme_applied_for_display()}")

    if work_experience:
        p.drawString(100, height - 500, f"Employer: {work_experience.employer}")
        p.drawString(100, height - 520, f"Work Details: {work_experience.work_details}")

    if documents:
        p.drawString(100, height - 540, f"Birth Certificate: {documents.birth_certificate.url if documents.birth_certificate else 'Not Provided'}")
        p.drawString(100, height - 560, f"ID Document: {documents.id_document.url if documents.id_document else 'Not Provided'}")

    p.showPage()
    p.save()
    buffer.seek(0)

    response = HttpResponse(buffer, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{application.first_name}_{application.last_name}.pdf"'
    return response


def download_word(request, application_id):
    application = get_object_or_404(Application, id=application_id)
    contact = application.contacts.first()  # Grab the first contact if available
    masters_edu = application.educational.first()  # First Masters educational record
    certificate_edu = application.certificate_edu.first()  # First Certificate educational record
    work_experience = application.work_experiences.first()  # First work experience

    # Create a Word document
    doc = Document()
    doc.add_heading(f'Application Detail for {application.first_name} {application.last_name}', level=1)

    doc.add_paragraph(f'Application Type: {application.get_application_type_display()}')
    doc.add_paragraph(f'First Name: {application.first_name}')
    doc.add_paragraph(f'Last Name: {application.last_name}')
    doc.add_paragraph(f'Date Of Birth: {application.date_of_birth}')
    doc.add_paragraph(f'Gender: {application.gender}')
    doc.add_paragraph(f'National ID: {application.national_id}')
    doc.add_paragraph(f'Title: {application.get_title_display()}')
    doc.add_paragraph(f'Marital Status: {application.get_marital_status_display()}')
    doc.add_paragraph(f'Place of Birth: {application.place_of_birth}')
    doc.add_paragraph(f'Citizenship: {application.citizenship}')
    doc.add_paragraph(f'Country: {application.country}')
    doc.add_paragraph(f'Disability: {application.disability}')
    doc.add_paragraph(f'Disability Details: {application.disability_details}')
    
    if contact:
        doc.add_paragraph(f'Email: {contact.email}')
        doc.add_paragraph(f'Phone Number: {contact.phone_number}')
        doc.add_paragraph(f'Permanent Address: {contact.permanent_address_no_street}')
        if contact.permanent_address_apt_unit:
            doc.add_paragraph(f'Apt/Unit: {contact.permanent_address_apt_unit}')
        doc.add_paragraph(f'State/Province: {contact.permanent_address_state_province}')

    if masters_edu:
        doc.add_paragraph(f'Programme Applied For (Masters): {masters_edu.get_programme_applied_for_display()}')
        doc.add_paragraph(f'Institution: {masters_edu.awarding_institution}')
        doc.add_paragraph(f'Major Subjects: {masters_edu.major_subjects}')

    if certificate_edu:
        doc.add_paragraph(f'Programme Applied For (Certificate): {certificate_edu.get_programme_applied_for_display()}')

    if work_experience:
        doc.add_paragraph(f'Employer: {work_experience.employer}')
        doc.add_paragraph(f'Work Details: {work_experience.work_details}')

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response = HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{application.first_name}_{application.last_name}.docx"'
    return response


def download_excel(request, application_id):
    application = get_object_or_404(Application, id=application_id)
    contact = application.contacts.first()  # First contact record
    masters_edu = application.educational.first()  # First Masters educational record
    certificate_edu = application.certificate_edu.first()  # First Certificate educational record
    work_experience = application.work_experiences.first()  # First work experience

    # Create a Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "Application Details"

    headers = [
        "First Name", 
        "Last Name", 
        "Date Of Birth", 
        "Gender", 
        "National ID", 
        "Title", 
        "Marital Status",
        "Place of Birth",
        "Citizenship",
        "Country",
        "Disability",
        "Disability Details",
        "Email",
        "Phone Number",
        "Permanent Address",
        "Apt/Unit",
        "State/Province",
        "Masters Programme",
        "Masters Institution",
        "Masters Major Subjects",
        "Certificate Programme",
        "Work Employer",
        "Work Details",
        "Documents (Birth Certificate)",
        "Documents (ID)"
    ]

    ws.append(headers)

    row = [
        application.first_name,
        application.last_name,
        application.date_of_birth,
        application.gender,
        application.national_id,
        application.get_title_display(),
        application.get_marital_status_display(),
        application.place_of_birth,
        application.citizenship,
        application.country,
        application.disability,
        application.disability_details,
        contact.email if contact else '',
        contact.phone_number if contact else '',
        contact.permanent_address_no_street if contact else '',
        contact.permanent_address_apt_unit if contact else '',
        contact.permanent_address_state_province if contact else '',
        masters_edu.get_programme_applied_for_display() if masters_edu else '',
        masters_edu.awarding_institution if masters_edu else '',
        masters_edu.major_subjects if masters_edu else '',
        certificate_edu.get_programme_applied_for_display() if certificate_edu else '',
        work_experience.employer if work_experience else '',
        work_experience.work_details if work_experience else '',
        application.documents.first().birth_certificate.url if application.documents.exists() and application.documents.first().birth_certificate else 'Not Provided',
        application.documents.first().id_document.url if application.documents.exists() and application.documents.first().id_document else 'Not Provided',
    ]

    ws.append(row)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename="{application.first_name}_{application.last_name}.xlsx"'

    wb.save(response)
    return response
    
