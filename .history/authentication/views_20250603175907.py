# import this to require login
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.contrib.auth import authenticate, login
# import this for sending email to user
from django.contrib.auth.tokens import default_token_generator
from django.shortcuts import render, redirect
from datetime import  date

from authentication.forms import UserRegistrationForm, UserLoginForm
from django.contrib.auth.models import Group
from .models import (
    Application, 
    Contact, 
    Documents, 
    WorkExperience, 
    Other, 
    Status,
    MastersEducational,
    CertificateEducational,
    ApplicationPayments,
    GeneralPayments,
)
from django.shortcuts import render, get_object_or_404
from django import forms
from django.urls import reverse
from django.conf import settings
from .forms import (
    CertificateEducationalDocumentForm,
    DocumentUploadForm,
    EducationalBackgroundCertificateForm,
    EducationalBackgroundMastersForm,
    MastersEducationalDocumentForm,
    ProgramTypeForm,
    StatusForm,
    UserLoginForm,
    UserRegistrationForm,
    PersonalInfoForm,
    ContactAndAddressForm,
    WorkExperienceForm,
    OtherInfoForm,
    
)

from paynow import Paynow 
from django.utils.http import urlsafe_base64_encode
from django.utils.encoding import force_bytes
from django.template.loader import render_to_string
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
from openpyxl import Workbook
from io import BytesIO
from django.http import HttpResponse
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt

import base64
import json
from django.core.files.base import ContentFile
from django.core.exceptions import ValidationError
import logging

logger = logging.getLogger(__name__)

# Create your views here.

from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt

from django.views.generic import TemplateView

class PrivacyPolicyView(TemplateView):
    template_name = 'privacy_policy.html'
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['last_updated'] = "April 30, 2025"  # Update this date when changing policy
        return context

def form_complete(request):
    # Render the success template
    return render(request, 'form_complete.html')

def payment_cancelled(request):
    # Render the cancellation template
    return render(request, 'payment_cancelled.html')

def application_list(request):
    applications = Application.objects.all()  # Fetch all applications
    return render(request, 'application/application_list.html', {'applications': applications})

def payment_timeout(request):
    return render(request, 'authentication/payment_mobile.html')

def application_detail(request, application_id):
    application = get_object_or_404(Application, id=application_id)
    return render(request, 'application/application_detail.html', {'application': application})

def cookie_policy(request):
    return render(request, 'cookie.html')



@login_required(login_url='login')
def homepage(request):
    user = request.user
    try:
        application = Application.objects.get(user=user)
        # fetch the most recent status
        latest_status = application.statuses.order_by('-created_at').first()
        status_text = latest_status.status_text if latest_status else None
    except Application.DoesNotExist:
        application = None
        status_text = None
    
    context = {
        'application': application,
        'status_text': status_text,
    }
    return render(request, 'homepage.html')


def register(request):
    if request.method == 'POST':
        form = UserRegistrationForm(request.POST)
        if form.is_valid():
            user = form.save()
            user.is_active = True  # Activate the user immediately
            user.save()

            # Assign user to the "Student" group
            try:
                student_group = Group.objects.get(name='Students')
                student_group.user_set.add(user)  # Add the user to the group
            except Group.DoesNotExist:
                # Handle the case if the group does not exist
                messages.error(request, 'User group not found. Please contact support.')

            login(request, user)  # Log the user in
            messages.success(request, 'Your account has been created and activated successfully!')
            return redirect('homepage')  # Redirect to the homepage or another page
        else:
            for field, errors in form.errors.items():
                for error in errors:
                    messages.error(request, f"{field}: {error}")
    else:
        form = UserRegistrationForm()
    return render(request, 'authentication/register.html', {'form': form})


def login_view(request):
    if request.method == 'POST':
        form = UserLoginForm(data=request.POST)
        if form.is_valid():
            username = form.cleaned_data.get('username')
            password = form.cleaned_data.get('password')
            user = authenticate(request, username=username, password=password)

            if user is not None:
                # Successful login
                login(request, user)  # Log the user in

                # Check user group and redirect accordingly
                if user.groups.filter(name='Admin').exists():
                    return redirect('/admin/')  # Redirect for Registrars
                elif user.groups.filter(name='Students').exists():
                    return redirect('homepage')  # Redirect for Students
                else:
                    # Redirect or message for users not in the specified groups
                    messages.warning(request, 'You do not have permission to access this application.')
                    return redirect('login')  # Redirect back to login or another page
            else:
                messages.error(request, 'Invalid username or password.')  # Invalid credentials
        else:
            # Handle form errors by sending feedback to the user
            for field, errors in form.errors.items():
                for error in errors:
                    messages.error(request, error)

    else:
        form = UserLoginForm()  # Initialize empty form for GET requests

    # Render the login template with the form
    return render(request, 'authentication/login.html', {'form': form})



from django.shortcuts import render, redirect
from django.contrib import messages
from .forms import (
    PersonalInfoForm,
    ContactAndAddressForm,
    EducationalBackgroundMastersForm,
    EducationalBackgroundCertificateForm,
    DocumentUploadForm,
    WorkExperienceForm,
    OtherInfoForm,
    
)

def check_payment_status(paynow, response, application):
    """Check the payment status using the poll URL."""
    if 'Test Case: Success' in response.instruction:
        application.is_paid = True  # Update the is_paid field
        application.save()  # Save the application
        return True
    
      # Handle live payments
    if response.success:
        status = paynow.check_transaction_status(response.poll_url)
        if status.paid:
            application.is_paid = True
            application.save()
            return True
    return False


def Paynow_Payment(request, application_id):
    application = get_object_or_404(Application, id=application_id)
    paynow = Paynow(
        settings.PAYNOW['INTEGRATION_ID'],
        settings.PAYNOW['INTEGRATION_KEY'],
        settings.PAYNOW['RESULT_URL'],
        settings.PAYNOW['RETURN_URL']
    )

    if request.method == 'POST':
        # Extract payment details directly from POST data
        name = request.POST.get('name')
        payment_type = request.POST.get('payment_type')
        amount = request.POST.get('amount')
        email = request.POST.get('email')
        phone = request.POST.get('phone')
        reference = request.POST.get('reference')
        payment_method = request.POST.get('payment_method')

        # Create payment with the user-provided details
        paynow_payment = paynow.create_payment(reference, email)
        paynow_payment.add(f"{payment_type} - {name}", float(amount))

        if payment_method == 'mobile':
            response = paynow.send_mobile(paynow_payment, phone, 'ecocash')
            print(f"Mobile Payment Response: {response.__dict__}")



            if response.success:

                is_test_case = 'TESTING: Faked Success' in response.instruction
                is_paid = is_test_case or paynow.check_transaction_status(response.poll_url)

                if is_paid:
                    application.is_paid = True
                    application.save()

                # Save Paynow response data to the database
                ApplicationPayments.objects.create(
                    application=application,
                    paynow_reference=response.data.get('paynowreference', ''),
                    status='paid' if 'Test Case: Success' in response.instruction else 'failed',
                    customer_email=email,
                    customer_name=name,  
                    customer_phone=phone,
                    amount=amount,
                    reference=reference
                )
                messages.success(request, 'Thank you for Application, your application is now processed')
                return redirect('payment_timeout')
            else:
                messages.error(request, 'Mobile payment failed. Please try again.')
                return redirect('paynow_payment', application_id=application.id)
              

        elif payment_method == 'card':
            response = paynow.send(paynow_payment)
            print(f"Card Payment Response: {response.__dict__}")

            if response.success:
                ApplicationPayments.objects.create(
                    application=application,
                    redirect_url=response.data.get('redirect_url', ''),
                    status='paid' if getattr(response, 'success', False) else 'failed',
                    customer_email=email,
                    customer_name=name,  
                    customer_phone=phone,
                    amount=amount,
                    reference=reference
                )
                
                if response.has_redirect:
                    # Redirect to Paynow for card payment
                    return redirect(response.redirect_url)
                else:
                    messages.error(request, 'Card payment didn\'t redirect. Please contact support.')
            else:
                messages.error(request, 'Card payment failed. Please try again.')
    else:
        # If it's a GET request, render the payment page without a form
        context = {
            'application': application,
        }
        return render(request, 'application/paynow_payment.html', context)
    


def convert_dates(data):
    """Convert date objects to strings in nested structures."""
    if isinstance(data, dict):
        return {k: v.isoformat() if isinstance(v, date) else convert_dates(v) 
                for k, v in data.items()}
    elif isinstance(data, (list, tuple)):
        return [convert_dates(item) for item in data]
    return data

def prepare_session_data(form_data):
    """Prepare form data for session storage by removing non-serializable fields."""
    if isinstance(form_data, dict):
        return {k: prepare_session_data(v) 
                for k, v in form_data.items() 
                if not hasattr(v, 'file') and not isinstance(v, (list, tuple))}  # Exclude file fields and nested structures
    elif isinstance(form_data, (list, tuple)):
        return [prepare_session_data(item) for item in form_data 
                if not hasattr(item, 'file')]  # Exclude file fields in lists

    return form_data

def prepare_uploaded_files(files):
    """Extract file information into a serializable format."""
    return {file.name: {'size': file.size, 'content_type': file.content_type} for file in files.values()}

@login_required(login_url='login')
def masters_application(request):
    if request.method == "POST":
        # Create all the forms
        personal_info_form = PersonalInfoForm(request.POST)
        documents_upload_form = DocumentUploadForm(request.POST, request.FILES)
        contact_and_address_form = ContactAndAddressForm(request.POST)
        educational_form = EducationalBackgroundMastersForm(request.POST)
        document_upload_form = MastersEducationalDocumentForm(request.POST, request.FILES)
        work_experience_form = WorkExperienceForm(request.POST)
        other_info_form = OtherInfoForm(request.POST)

        # Validate all forms
        if (personal_info_form.is_valid() and documents_upload_form.is_valid() and 
            contact_and_address_form.is_valid() and educational_form.is_valid() and 
            document_upload_form.is_valid() and work_experience_form.is_valid() and 
            other_info_form.is_valid()):
            
            print("All forms are valid. Saving data...")
            
            # Create and save the Application first
            application = Application.objects.create(
                user=request.user,
                application_type='masters',
                **personal_info_form.cleaned_data
            )
            print("Application saved:", application.id)

            # Save documents and associate with the application
            documents = documents_upload_form.save(commit=False)
            documents.application = application  # Associate with the application
            documents.save()
            print("Documents saved:", documents.id)

            # Save other related models
            Contact.objects.create(
                application=application,
                **contact_and_address_form.cleaned_data
            )
            print("Contact saved")

            MastersEducational.objects.create(
                application=application,
                **educational_form.cleaned_data
            )
            print("MastersEducational saved")

            # Save documents and associate with the application
            document = document_upload_form.save(commit=False)
            document.application = application  # Associate with the application
            document.save()
            print("MastersDocuments saved:", document.id)

            WorkExperience.objects.create(
                application=application,
                **work_experience_form.cleaned_data
            )
            print("WorkExperience saved")

            Other.objects.create(
                application=application,
                **other_info_form.cleaned_data
            )
            print("Other saved")

            Status.objects.create(
                application=application,
                status='Submitted'
            )
            print("Status saved")

            # Redirect to payment  page
            return redirect('paynow_payment',application_id=application.id)

        else:
            print("Form errors:")
            print("Personal Info:", personal_info_form.errors)
            print("Documents Upload:", documents_upload_form.errors)
            print("Contact and Address:", contact_and_address_form.errors)
            print("Educational:", educational_form.errors)
            print("Document Upload:", document_upload_form.errors)
            print("Work Experience:", work_experience_form.errors)
            print("Other Info:", other_info_form.errors)
            messages.error(request, 'Please correct the errors in your application data.')
            return render(request, 'application/masters_application.html', {
                'personal_info_form': personal_info_form,
                'documents_upload_form': documents_upload_form,
                'contact_and_address_form': contact_and_address_form,
                'educational_form': educational_form,
                'document_upload_form': document_upload_form,
                'work_experience_form': work_experience_form,
                'other_info_form': other_info_form,
            })

    else:
        # Initialize the forms
        personal_info_form = PersonalInfoForm()
        documents_upload_form = DocumentUploadForm()
        contact_and_address_form = ContactAndAddressForm()
        educational_form = EducationalBackgroundMastersForm()
        document_upload_form = MastersEducationalDocumentForm()
        work_experience_form = WorkExperienceForm()
        other_info_form = OtherInfoForm()

    return render(request, 'application/masters_application.html', {
        'personal_info_form': personal_info_form,
        'documents_upload_form': documents_upload_form,
        'contact_and_address_form': contact_and_address_form,
        'educational_form': educational_form,
        'document_upload_form': document_upload_form,
        'work_experience_form': work_experience_form,
        'other_info_form': other_info_form,
    })

@login_required(login_url='login')
def certificate_application(request):
    if request.method == "POST":
        personal_info_form = PersonalInfoForm(request.POST)
        document_upload_form = DocumentUploadForm(request.POST, request.FILES)
        contact_and_address_form = ContactAndAddressForm(request.POST)
        educational_form = EducationalBackgroundCertificateForm(request.POST)
        documents_upload_form = CertificateEducationalDocumentForm(request.POST, request.FILES)
        other_info_form = OtherInfoForm(request.POST)

        if (personal_info_form.is_valid() and document_upload_form.is_valid() and contact_and_address_form.is_valid() and
            educational_form.is_valid() and documents_upload_form.is_valid() and
            other_info_form.is_valid()):

            # Create and save the application first
            application = Application.objects.create(
                user=request.user,
                application_type='certificate',
                **personal_info_form.cleaned_data
            )
            
            # Save documents and associate with the application
            document = document_upload_form.save(commit=False)
            document.application = application  # Associate with the application
            document.save()
            
            # Save other related models
            Contact.objects.create(
                application=application,
                **contact_and_address_form.cleaned_data
            )

            CertificateEducational.objects.create(
                application=application,
                **educational_form.cleaned_data
            )
            
            # Save documents and associate with the application
            documents = documents_upload_form.save(commit=False)
            documents.application = application  # Associate with the application
            documents.save()
            
            Other.objects.create(
                application=application,
                **other_info_form.cleaned_data
            )
            Status.objects.create(
                application=application,
                status='Submitted'
            )
            
            # Redirect to payment page with the application_id
            return redirect('paynow_payment', application_id=application.id)  # Pass application_id

        else:
            messages.error(request, 'Please correct the errors in your application data.')

    else:
        personal_info_form = PersonalInfoForm()
        document_upload_form = DocumentUploadForm()
        contact_and_address_form = ContactAndAddressForm()
        educational_form = EducationalBackgroundCertificateForm()
        documents_upload_form = CertificateEducationalDocumentForm()
        other_info_form = OtherInfoForm()

    return render(request, 'application/certificate_application.html', {
        'personal_info_form': personal_info_form,
        'document_upload_form': document_upload_form,
        'contact_and_address_form': contact_and_address_form,
        'educational_form': educational_form,
        'documents_upload_form': documents_upload_form,
        'other_info_form': other_info_form,
    })
    
def track_application(request):
    application = None
    error_message = None
    
    if request.method == 'POST':
        application_id = request.POST.get('application_id')
        
        try:
            # Convert to integer and find application
            application = Application.objects.get(id=int(application_id))
            latest_status = application.statuses.order_by('-id').first()
        except (ValueError, Application.DoesNotExist):
            error_message = "Invalid application ID. Please check your tracking number."
        except Exception as e:
            error_message = f"Error retrieving application: {str(e)}"

    return render(request, 'application/track_application.html', {
        'application': application,
        'latest_status': application.statuses.order_by('-id').first() if application else None,
        'error_message': error_message
    })
    
    
@login_required 
def update_status(request, application_id):
    application = get_object_or_404(Application, id=application_id)
    
    
    status, created = Status.objects.get_or_create(application=application)
    
    if request.method == 'POST':
        form = StatusForm(request.POST, instance=status)
        if form.is_valid():
            form.save()  
            return redirect('application_detail', application_id=application.id) 
    else:
        form = StatusForm(instance=status)

    return render(request, 'application/update_status.html', {'form': form, 'application': application})
   
def registrar_view(request, application_id):
    application = get_object_or_404(Application, id=application_id)
    
    status_instance = Status.objects.filter(application=application).order_by('-updated_at').first()

    if status_instance is None:
        
        status_instance = Status(application=application, status='Pending')
        status_instance.save()

    if request.method == "POST":
        form = StatusForm(request.POST, instance=status_instance)
        if form.is_valid():
            form.save() 
            return redirect('application_detail', application_id=application_id)  
    else:
        form = StatusForm(instance=status_instance)

    
    return render(request, 'application/application_detail.html', {
        'form': form,
        'application': application,
        'status': status_instance.status,  
        'last_update': status_instance.last_update  
    }) 
    
    
   
def download_pdf(request, application_id):
    application = get_object_or_404(Application, id=application_id)
    contact = application.contacts.first()  # Assuming there's at least one contact related
    masters_edu = application.educational.first()  # Assuming the first educational record for masters
    certificate_edu = application.certificate_edu.first()  # Assuming the first educational record for certificate
    work_experience = application.work_experiences.first()  # Fetch first work experience for simplicity
    documents = application.documents.first()  # Assuming one set of documents is related

    # Create a PDF document
    buffer = BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    
    p.drawString(100, height - 40, f"Application Detail for {application.first_name} {application.last_name}")
    p.drawString(100, height - 60, f"Application Type: {application.get_application_type_display()}")
    p.drawString(100, height - 80, f"First Name: {application.first_name}")
    p.drawString(100, height - 100, f"Last Name: {application.last_name}")
    p.drawString(100, height - 120, f"Date Of Birth: {application.date_of_birth}")
    p.drawString(100, height - 140, f"Gender: {application.gender}")
    p.drawString(100, height - 160, f"National ID: {application.national_id}")
    p.drawString(100, height - 180, f"Title: {application.get_title_display()}")
    p.drawString(100, height - 200, f"Marital Status: {application.get_marital_status_display()}")
    p.drawString(100, height - 220, f"Place of Birth: {application.place_of_birth}")
    p.drawString(100, height - 240, f"Citizenship: {application.citizenship}")
    p.drawString(100, height - 260, f"Country: {application.country}")
    p.drawString(100, height - 280, f"Disability: {application.disability}")
    p.drawString(100, height - 300, f"Disability Details: {application.disability_details}")
    
    if contact:
        p.drawString(100, height - 320, f"Email: {contact.email}")
        p.drawString(100, height - 340, f"Phone Number: {contact.phone_number}")
        p.drawString(100, height - 360, f"Permanent Address: {contact.permanent_address_no_street}")
        if contact.permanent_address_apt_unit:
            p.drawString(100, height - 380, f"Apt/Unit: {contact.permanent_address_apt_unit}")
        p.drawString(100, height - 400, f"State/Province: {contact.permanent_address_state_province}")

    if masters_edu:
        p.drawString(100, height - 420, f"Programme Applied For (Masters): {masters_edu.get_programme_applied_for_display()}")
        p.drawString(100, height - 440, f"Institution: {masters_edu.awarding_institution}")
        p.drawString(100, height - 460, f"Major Subjects: {masters_edu.major_subjects}")
    
    if certificate_edu:
        p.drawString(100, height - 480, f"Programme Applied For (Certificate): {certificate_edu.get_programme_applied_for_display()}")

    if work_experience:
        p.drawString(100, height - 500, f"Employer: {work_experience.employer}")
        p.drawString(100, height - 520, f"Work Details: {work_experience.work_details}")

    if documents:
        p.drawString(100, height - 540, f"Birth Certificate: {documents.birth_certificate.url if documents.birth_certificate else 'Not Provided'}")
        p.drawString(100, height - 560, f"ID Document: {documents.id_document.url if documents.id_document else 'Not Provided'}")

    p.showPage()
    p.save()
    buffer.seek(0)

    response = HttpResponse(buffer, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{application.first_name}_{application.last_name}.pdf"'
    return response


def download_word(request, application_id):
    application = get_object_or_404(Application, id=application_id)
    contact = application.contacts.first()  # Grab the first contact if available
    masters_edu = application.educational.first()  # First Masters educational record
    certificate_edu = application.certificate_edu.first()  # First Certificate educational record
    work_experience = application.work_experiences.first()  # First work experience

    # Create a Word document
    doc = Document()
    doc.add_heading(f'Application Detail for {application.first_name} {application.last_name}', level=1)

    doc.add_paragraph(f'Application Type: {application.get_application_type_display()}')
    doc.add_paragraph(f'First Name: {application.first_name}')
    doc.add_paragraph(f'Last Name: {application.last_name}')
    doc.add_paragraph(f'Date Of Birth: {application.date_of_birth}')
    doc.add_paragraph(f'Gender: {application.gender}')
    doc.add_paragraph(f'National ID: {application.national_id}')
    doc.add_paragraph(f'Title: {application.get_title_display()}')
    doc.add_paragraph(f'Marital Status: {application.get_marital_status_display()}')
    doc.add_paragraph(f'Place of Birth: {application.place_of_birth}')
    doc.add_paragraph(f'Citizenship: {application.citizenship}')
    doc.add_paragraph(f'Country: {application.country}')
    doc.add_paragraph(f'Disability: {application.disability}')
    doc.add_paragraph(f'Disability Details: {application.disability_details}')
    
    if contact:
        doc.add_paragraph(f'Email: {contact.email}')
        doc.add_paragraph(f'Phone Number: {contact.phone_number}')
        doc.add_paragraph(f'Permanent Address: {contact.permanent_address_no_street}')
        if contact.permanent_address_apt_unit:
            doc.add_paragraph(f'Apt/Unit: {contact.permanent_address_apt_unit}')
        doc.add_paragraph(f'State/Province: {contact.permanent_address_state_province}')

    if masters_edu:
        doc.add_paragraph(f'Programme Applied For (Masters): {masters_edu.get_programme_applied_for_display()}')
        doc.add_paragraph(f'Institution: {masters_edu.awarding_institution}')
        doc.add_paragraph(f'Major Subjects: {masters_edu.major_subjects}')

    if certificate_edu:
        doc.add_paragraph(f'Programme Applied For (Certificate): {certificate_edu.get_programme_applied_for_display()}')

    if work_experience:
        doc.add_paragraph(f'Employer: {work_experience.employer}')
        doc.add_paragraph(f'Work Details: {work_experience.work_details}')

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response = HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{application.first_name}_{application.last_name}.docx"'
    return response


def download_excel(request, application_id):
    application = get_object_or_404(Application, id=application_id)
    contact = application.contacts.first()  # First contact record
    masters_edu = application.educational.first()  # First Masters educational record
    certificate_edu = application.certificate_edu.first()  # First Certificate educational record
    work_experience = application.work_experiences.first()  # First work experience

    # Create a Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "Application Details"

    headers = [
        "First Name", 
        "Last Name", 
        "Date Of Birth", 
        "Gender", 
        "National ID", 
        "Title", 
        "Marital Status",
        "Place of Birth",
        "Citizenship",
        "Country",
        "Disability",
        "Disability Details",
        "Email",
        "Phone Number",
        "Permanent Address",
        "Apt/Unit",
        "State/Province",
        "Masters Programme",
        "Masters Institution",
        "Masters Major Subjects",
        "Certificate Programme",
        "Work Employer",
        "Work Details",
        "Documents (Birth Certificate)",
        "Documents (ID)"
    ]

    ws.append(headers)

    row = [
        application.first_name,
        application.last_name,
        application.date_of_birth,
        application.gender,
        application.national_id,
        application.get_title_display(),
        application.get_marital_status_display(),
        application.place_of_birth,
        application.citizenship,
        application.country,
        application.disability,
        application.disability_details,
        contact.email if contact else '',
        contact.phone_number if contact else '',
        contact.permanent_address_no_street if contact else '',
        contact.permanent_address_apt_unit if contact else '',
        contact.permanent_address_state_province if contact else '',
        masters_edu.get_programme_applied_for_display() if masters_edu else '',
        masters_edu.awarding_institution if masters_edu else '',
        masters_edu.major_subjects if masters_edu else '',
        certificate_edu.get_programme_applied_for_display() if certificate_edu else '',
        work_experience.employer if work_experience else '',
        work_experience.work_details if work_experience else '',
        application.documents.first().birth_certificate.url if application.documents.exists() and application.documents.first().birth_certificate else 'Not Provided',
        application.documents.first().id_document.url if application.documents.exists() and application.documents.first().id_document else 'Not Provided',
    ]

    ws.append(row)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename="{application.first_name}_{application.last_name}.xlsx"'

    wb.save(response)
    return response


    
def validate_file(file):
    # Max size: 5MB
    max_size = 5 * 1024 * 1024
    if file.size > max_size:
        raise ValidationError("File size exceeds 5MB limit")
    
    # Allowed file types
    allowed_types = ['application/pdf', 'image/jpeg', 'image/png']
    if file.content_type not in allowed_types:
        raise ValidationError("Only PDF, JPEG, and PNG files are allowed")
    


        


def general_payment(request):
    paynow = Paynow(
        settings.PAYNOW['INTEGRATION_ID'],
        settings.PAYNOW['INTEGRATION_KEY'],
        settings.PAYNOW['RESULT_URL'],
        settings.PAYNOW['RETURN_URL']
    )

    if request.method == 'POST':
        # Extract payment details directly from POST data
        name = request.POST.get('name')
        payment_type = request.POST.get('payment_type')
        amount = request.POST.get('amount')
        email = request.POST.get('email')
        phone = request.POST.get('phone')
        reference = request.POST.get('reference')
        payment_method = request.POST.get('payment_method')

        # Create payment with the user-provided details
        paynow_payment = paynow.create_payment(reference, email)
        paynow_payment.add(f"{payment_type} - {name}", float(amount))

        if payment_method == 'mobile':
            response = paynow.send_mobile(paynow_payment, phone, 'ecocash')
            print(f"Mobile Payment Response: {response.__dict__}")



            if response.success:
                
                # Save Paynow response data to the database
                GeneralPayments.objects.create(
                    paynow_reference=response.data.get('paynowreference', ''),
                    customer_email=email,
                    customer_name=name,  
                    customer_phone=phone,
                    amount=amount,
                    reference=reference
                )
                messages.success(request, 'Payment successful. Payment data saved. Thank you for your payment')
                return redirect('payment_timeout')
            else:
                messages.error(request, 'Mobile payment failed. Please try again.')
                return redirect('general_payment')
               
        elif payment_method == 'card':
            response = paynow.send(paynow_payment)
            print(f"Card Payment Response: {response.__dict__}")

            if response.success:
                redirect_url = response.data.get('redirect_url')
                paynow_reference = ''
                if redirect_url:
                    print("Redirect URL:", redirect_url)  # Debug
        # Remove trailing slash if any
                    url = redirect_url.rstrip('/')
        # Extract the last segment (payment reference)
                    paynow_reference = url.split('/')[-1]
                    print("Extracted paynow_reference:", paynow_reference) 
                GeneralPayments.objects.create(
                    paynow_reference=paynow_reference,  
                    redirect_url=response.data.get('redirect_url', ''),
                    status='paid' if getattr(response, 'success', False) else 'failed',
                    customer_email=email,
                    customer_name=name,  
                    customer_phone=phone,
                    amount=amount,
                    reference=reference
                )
                if response.has_redirect:
                    # Redirect to Paynow for card payment
                    return redirect(response.redirect_url)
                else:
                    messages.error(request, 'Card payment didn\'t redirect. Please contact support.')
                    return redirect('general_payment')
            else:
                messages.error(request, 'Card payment failed. Please try again.')
                return redirect('general_payment')

    else:
       
        return render(request, 'authentication/general_payment.html')
    


def card_payment(request, application_id):
    application = get_object_or_404(Application, id=application_id)
    paynow = Paynow(
        settings.PAYNOW['INTEGRATION_ID'],
        settings.PAYNOW['INTEGRATION_KEY'],
        settings.PAYNOW['RESULT_URL'],
        settings.PAYNOW['RETURN_URL']
    )

    if request.method == 'POST':
        # Extract payment details directly from POST data
        name = request.POST.get('name')
        payment_type = request.POST.get('payment_type')
        amount = request.POST.get('amount')
        email = request.POST.get('email')
        phone = request.POST.get('phone')
        reference = request.POST.get('reference')
        payment_method = request.POST.get('payment_method')

        # Create payment with the user-provided details
        paynow_payment = paynow.create_payment(reference, email)
        paynow_payment.add(f"{payment_type} - {name}", float(amount))

        if payment_method == 'mobile':
            response = paynow.send_mobile(paynow_payment, phone, 'ecocash')
            print(f"Mobile Payment Response: {response.__dict__}")



            if response.success:

                is_test_case = 'TESTING: Faked Success' in response.instruction
                is_paid = is_test_case or paynow.check_transaction_status(response.poll_url)

                if is_paid:
                    application.is_paid = True
                    application.save()

                # Save Paynow response data to the database
                ApplicationPayments.objects.create(
                    application=application,
                    paynow_reference=response.data.get('paynowreference', ''),
                    status='paid' if 'Test Case: Success' in response.instruction else 'failed',
                    customer_email=email,
                    customer_name=name,  
                    customer_phone=phone,
                    amount=amount,
                    reference=reference
                )
                messages.success(request, 'Thank you for Application, your application is now processed')
                return redirect('payment_timeout')
            else:
                messages.error(request, 'Mobile payment failed. Please try again.')
                return redirect('paynow_payment', application_id=application.id)
              

        elif payment_method == 'card':
            response = paynow.send(paynow_payment)
            print(f"Card Payment Response: {response.__dict__}")

            if response.success:
                ApplicationPayments.objects.create(
                    application=application,
                    redirect_url=response.data.get('redirect_url', ''),
                    status='paid' if getattr(response, 'success', False) else 'failed',
                    customer_email=email,
                    customer_name=name,  
                    customer_phone=phone,
                    amount=amount,
                    reference=reference
                )
                
                if response.has_redirect:
                    # Redirect to Paynow for card payment
                    return redirect(response.redirect_url)
                else:
                    messages.error(request, 'Card payment didn\'t redirect. Please contact support.')
            else:
                messages.error(request, 'Card payment failed. Please try again.')
    else:
        # If it's a GET request, render the payment page without a form
        context = {
            'application': application,
        }
        return render(request, 'application/card_payment.html', context)
    

    
